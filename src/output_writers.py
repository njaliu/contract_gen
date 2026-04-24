import json
import random
import re
from pathlib import Path
from typing import Any

def _align_from_text(value: str | None):
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if not value:
        return None
    mapping = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    return mapping.get(value.lower())


def _table_align_from_text(value: str | None):
    from docx.enum.table import WD_TABLE_ALIGNMENT

    if not value:
        return None
    mapping = {
        "left": WD_TABLE_ALIGNMENT.LEFT,
        "center": WD_TABLE_ALIGNMENT.CENTER,
        "right": WD_TABLE_ALIGNMENT.RIGHT,
    }
    return mapping.get(value.lower())


def _cell_v_align_from_text(value: str | None):
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

    if not value:
        return None
    mapping = {
        "top": WD_CELL_VERTICAL_ALIGNMENT.TOP,
        "center": WD_CELL_VERTICAL_ALIGNMENT.CENTER,
        "bottom": WD_CELL_VERTICAL_ALIGNMENT.BOTTOM,
    }
    return mapping.get(value.lower())


def _apply_paragraph_style(paragraph, style_cfg: dict[str, Any], text: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt

    run = paragraph.add_run(text)
    font_name = style_cfg.get("font")
    if font_name:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if "size_pt" in style_cfg:
        run.font.size = Pt(float(style_cfg["size_pt"]))
    run.bold = bool(style_cfg.get("bold", False))

    align = _align_from_text(style_cfg.get("align"))
    if align is not None:
        paragraph.alignment = align
    if "line_spacing" in style_cfg:
        paragraph.paragraph_format.line_spacing = float(style_cfg["line_spacing"])
    if "first_line_indent_pt" in style_cfg:
        paragraph.paragraph_format.first_line_indent = Pt(
            float(style_cfg["first_line_indent_pt"])
        )
    if "space_before_pt" in style_cfg:
        paragraph.paragraph_format.space_before = Pt(float(style_cfg["space_before_pt"]))
    if "space_after_pt" in style_cfg:
        paragraph.paragraph_format.space_after = Pt(float(style_cfg["space_after_pt"]))


def _set_cell_border(cell, style: str = "single", width_pt: float = 0.5) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    size_val = str(max(2, int(width_pt * 8)))
    for edge in ("top", "left", "bottom", "right"):
        edge_tag = f"w:{edge}"
        element = borders.find(qn(edge_tag))
        if element is None:
            element = OxmlElement(edge_tag)
            borders.append(element)
        element.set(qn("w:val"), style)
        element.set(qn("w:sz"), size_val)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def _parse_markdown_table(lines: list[str], start_idx: int) -> tuple[list[list[str]], int]:
    table_lines: list[str] = []
    i = start_idx
    while i < len(lines):
        line = lines[i].strip()
        if "|" not in line:
            break
        table_lines.append(line)
        i += 1
    if len(table_lines) < 2:
        return [], start_idx
    sep = table_lines[1]
    if not re.match(r"^\|\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?$", sep):
        return [], start_idx

    rows: list[list[str]] = []
    for raw in [table_lines[0]] + table_lines[2:]:
        parts = [p.strip() for p in raw.strip("|").split("|")]
        rows.append(parts)
    return rows, i


def _write_docx(content: str, output_path: Path, format_spec: dict[str, Any]) -> None:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as e:
        raise RuntimeError(
            "输出 docx 需要 python-docx。请先执行: python3 -m pip install -r requirements.txt"
        ) from e

    doc = Document()
    section = doc.sections[0]
    page_cfg = format_spec.get("page", {})
    section.page_width = Pt(float(page_cfg.get("width_pt", 595.3)))
    section.page_height = Pt(float(page_cfg.get("height_pt", 841.9)))
    section.left_margin = Pt(float(page_cfg.get("margin_left_pt", 90.0)))
    section.right_margin = Pt(float(page_cfg.get("margin_right_pt", 90.0)))
    section.top_margin = Pt(float(page_cfg.get("margin_top_pt", 72.0)))
    section.bottom_margin = Pt(float(page_cfg.get("margin_bottom_pt", 72.0)))
    section.header_distance = Pt(float(page_cfg.get("header_distance_pt", 42.55)))
    section.footer_distance = Pt(float(page_cfg.get("footer_distance_pt", 49.6)))

    styles = format_spec.get("styles", {})
    table_cfg = format_spec.get("table", {})
    lines = content.splitlines()

    def classify_style_key(raw_line: str) -> tuple[str, str]:
        stripped_line = raw_line.strip()
        # First honor explicit markdown heading markers when present.
        if stripped_line.startswith("# "):
            return "title", stripped_line[2:].strip()
        if stripped_line.startswith("### "):
            return "h1", stripped_line[4:].strip()
        if stripped_line.startswith("#### "):
            return "h2", stripped_line[5:].strip()
        if stripped_line.startswith("##### "):
            return "h3", stripped_line[6:].strip()

        # Then apply docx-derived hierarchy rules as fallback.
        if re.match(r"^附件[:：]\s*", stripped_line):
            return "h2_emphasis", stripped_line
        if stripped_line == "验收结论":
            return "h3", stripped_line
        if re.match(r"^6\.[12]\s+", stripped_line):
            return "h2", stripped_line
        if re.match(r"^(?:[1-9]|1[0-2])\s+", stripped_line):
            return "h1", stripped_line
        if re.match(r"^\d+\.\d+(?:\.\d+)?\s+", stripped_line):
            return "body", stripped_line
        return "body", raw_line

    idx = 0
    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            idx += 1
            continue

        table_rows, consumed_idx = _parse_markdown_table(lines, idx)
        if table_rows:
            cols = max(len(r) for r in table_rows)
            table = doc.add_table(rows=len(table_rows), cols=cols)
            table.style = "Table Grid"
            table_align = _table_align_from_text(table_cfg.get("default_alignment"))
            if table_align is not None:
                table.alignment = table_align
            v_align = _cell_v_align_from_text(table_cfg.get("cell_vertical_align"))
            border_style = table_cfg.get("border_style", "single")
            border_width = float(table_cfg.get("border_width_pt", 0.5))

            for r_i, row in enumerate(table_rows):
                for c_i in range(cols):
                    cell_text = row[c_i] if c_i < len(row) else ""
                    cell = table.cell(r_i, c_i)
                    if v_align is not None:
                        cell.vertical_alignment = v_align
                    _set_cell_border(cell, style=border_style, width_pt=border_width)
                    p = cell.paragraphs[0]
                    style_key = "table_header" if r_i == 0 else "table_body"
                    _apply_paragraph_style(p, styles.get(style_key, {}), cell_text)
            idx = consumed_idx
            continue

        style_key, text = classify_style_key(line)

        paragraph = doc.add_paragraph()
        _apply_paragraph_style(paragraph, styles.get(style_key, {}), text)
        idx += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def _write_markdown(content: str, output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(content + "\n", encoding="utf-8")


def _load_format_spec(
    format_spec_dir: str, format_id: str | None = None, format_seed: int | None = None
) -> dict[str, Any]:
    spec_dir = Path(format_spec_dir)
    if not spec_dir.exists():
        raise FileNotFoundError(f"FormatSpec 目录不存在: {spec_dir}")

    if format_id:
        target = spec_dir / f"{format_id}.json"
        if not target.exists():
            raise FileNotFoundError(f"未找到指定 FormatSpec: {target}")
        return json.loads(target.read_text(encoding="utf-8"))

    candidates = sorted(spec_dir.glob("docx_format_*.json"))
    if not candidates:
        raise RuntimeError(f"FormatSpec 目录下未找到 docx_format_*.json: {spec_dir}")
    rnd = random.Random(format_seed)
    selected = rnd.choice(candidates)
    return json.loads(selected.read_text(encoding="utf-8"))


def write_contract_output(
    contract_text: str,
    output_file: str,
    output_format: str = "md",
    format_spec_dir: str = "resources/formats",
    format_id: str | None = None,
    format_seed: int | None = None,
) -> list[Path]:
    output_path = Path(output_file)
    fmt = output_format.lower()
    written: list[Path] = []

    if fmt in ("md", "both"):
        md_path = output_path.with_suffix(".md") if fmt == "both" else output_path
        _write_markdown(contract_text, md_path)
        written.append(md_path.resolve())

    if fmt in ("docx", "both"):
        docx_path = output_path.with_suffix(".docx")
        spec = _load_format_spec(
            format_spec_dir=format_spec_dir,
            format_id=format_id,
            format_seed=format_seed,
        )
        _write_docx(contract_text, docx_path, spec)
        written.append(docx_path.resolve())

    if fmt not in ("md", "docx", "both"):
        raise ValueError(f"不支持的 output_format: {output_format}")

    return written
